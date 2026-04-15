"""
_doc_sections_a.py

Document sections 00-06: cover, executive summary, problem, what is codi,
core concepts, flag system, presets.
"""

from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from _doc_helpers import (
    RL3_ACCENT,
    RL3_BLACK,
    RL3_GRAY,
    RL3_MID_GRAY,
    add_body,
    add_bullet,
    add_callout,
    add_code_block,
    add_heading,
    add_horizontal_rule,
    add_section_label,
    add_table,
)


def build_cover(doc: Document) -> None:
    cover_logo = doc.add_paragraph()
    cover_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cover_logo.paragraph_format.space_before = Pt(60)
    cover_logo.paragraph_format.space_after = Pt(0)
    rl_run = cover_logo.add_run("RL")
    rl_run.font.name = "Calibri"
    rl_run.font.size = Pt(64)
    rl_run.font.bold = True
    rl_run.font.color.rgb = RL3_BLACK
    three_run = cover_logo.add_run("3")
    three_run.font.name = "Calibri"
    three_run.font.size = Pt(64)
    three_run.font.bold = True
    three_run.font.color.rgb = RL3_ACCENT

    agency = doc.add_paragraph()
    agency.paragraph_format.space_before = Pt(0)
    agency.paragraph_format.space_after = Pt(30)
    ar = agency.add_run("AI AGENCY")
    ar.font.name = "Courier New"
    ar.font.size = Pt(10)
    ar.font.color.rgb = RL3_GRAY

    add_horizontal_rule(doc, RL3_ACCENT, 8)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(30)
    title_p.paragraph_format.space_after = Pt(8)
    tr = title_p.add_run("Codi")
    tr.font.name = "Calibri"
    tr.font.size = Pt(48)
    tr.font.bold = True
    tr.font.color.rgb = RL3_BLACK

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(8)
    sr = sub_p.add_run("Unified AI Agent Configuration Platform")
    sr.font.name = "Calibri"
    sr.font.size = Pt(18)
    sr.font.italic = True
    sr.font.color.rgb = RL3_GRAY

    tag_p = doc.add_paragraph()
    tag_p.paragraph_format.space_before = Pt(4)
    tag_p.paragraph_format.space_after = Pt(40)
    tagr = tag_p.add_run("One config. Every AI agent. Zero drift.")
    tagr.font.name = "Courier New"
    tagr.font.size = Pt(11)
    tagr.font.color.rgb = RL3_ACCENT

    add_horizontal_rule(doc, RL3_MID_GRAY, 4)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(16)
    mr = meta_p.add_run(f"Version 2.0  ·  {datetime.now().strftime('%B %Y')}")
    mr.font.name = "Calibri"
    mr.font.size = Pt(10)
    mr.font.color.rgb = RL3_GRAY

    doc.add_page_break()


def build_executive_summary(doc: Document) -> None:
    add_section_label(doc, "01", "Executive Summary")
    add_heading(doc, "What Is Codi?")
    add_body(
        doc,
        (
            "Codi is the single source of truth for AI agent configuration. "
            "It solves the configuration drift problem that every team faces when "
            "using more than one AI coding agent: Claude Code, Cursor, Codex, "
            "Windsurf, and Cline each read different file formats, have different "
            "capabilities, and require separate maintenance. "
            "Codi eliminates that duplication."
        ),
    )
    add_body(
        doc,
        (
            "Write your rules, skills, agents, and behavioral policies once in a "
            "single .codi/ directory. Run codi generate. "
            "Every agent receives a perfectly formatted, verified configuration file "
            "derived from the same source. One change propagates everywhere."
        ),
    )
    add_callout(
        doc,
        "Core value proposition",
        "Teams using Codi report zero configuration drift across agents, "
        "consistent policy enforcement (security, testing, code style), "
        "and a single PR review process for all AI agent configuration changes.",
    )
    add_heading(doc, "Key Numbers", level=2)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Built-in rule templates", "27"],
            ["Built-in skill templates", "44"],
            ["Built-in agent templates", "22"],
            ["Built-in command templates", "17"],
            ["Supported AI agents", "5"],
            ["Built-in presets", "6"],
            ["Behavioral flags", "16"],
            ["CLI commands", "19"],
            ["Current version", "2.0.0"],
        ],
        col_widths=[3.5, 3.0],
    )


def build_problem(doc: Document) -> None:
    add_section_label(doc, "02", "The Problem")
    add_heading(doc, "Configuration Drift Across AI Agents")
    add_body(
        doc,
        (
            "Modern engineering teams no longer use a single AI coding tool. "
            "They use Claude Code for complex reasoning, Cursor for in-editor "
            "autocomplete, Codex for CLI automation, and Windsurf or Cline for "
            "additional context. Each agent reads a different configuration format:"
        ),
    )
    add_table(
        doc,
        ["Agent", "Config File", "Rules", "Skills", "Agents"],
        [
            [
                "Claude Code",
                "CLAUDE.md",
                ".claude/rules/",
                ".claude/skills/",
                ".claude/agents/",
            ],
            ["Cursor", ".cursorrules", ".cursor/rules/", ".cursor/skills/", "—"],
            ["Codex", "AGENTS.md", "inline", ".agents/skills/", ".codex/agents/"],
            ["Windsurf", ".windsurfrules", "inline", ".windsurf/skills/", "—"],
            ["Cline", ".clinerules", "inline", ".cline/skills/", "—"],
        ],
        col_widths=[1.3, 1.5, 1.5, 1.5, 1.2],
    )
    add_body(
        doc,
        (
            "A security rule added to CLAUDE.md never reaches .cursorrules. "
            "A team workflow skill defined for Cursor is invisible to Codex. "
            "When a teammate manually edits a generated file, the change silently "
            "diverges from the canonical policy. There is no single place to review, "
            "approve, or enforce AI agent configuration. That is the problem Codi solves."
        ),
    )
    add_heading(doc, "Why This Matters", level=2)
    add_bullet(
        doc,
        "Security rules defined in one agent never propagate to others — gaps remain.",
        "Security gaps.",
    )
    add_bullet(
        doc,
        "Code style and testing policies drift between tools — inconsistent enforcement.",
        "Policy drift.",
    )
    add_bullet(
        doc,
        "Teams maintain 5 separate config files with duplicate content — wasted effort.",
        "Maintenance burden.",
    )
    add_bullet(
        doc,
        "No audit trail — no record of who changed what and when.",
        "No visibility.",
    )
    add_bullet(
        doc,
        "Onboarding new agents requires recreating configuration from scratch.",
        "Slow onboarding.",
    )


def build_what_is_codi(doc: Document) -> None:
    add_section_label(doc, "03", "What Is Codi")
    add_heading(doc, "The Unified Configuration Platform")
    add_body(
        doc,
        (
            "Codi is a CLI tool that reads a single .codi/ directory "
            "and generates native configuration files for every AI agent your team uses. "
            "It is not a middleware layer, a proxy, or a runtime wrapper — it is a "
            "build-time configuration compiler."
        ),
    )
    add_callout(
        doc,
        "Design principle",
        "Codi operates at configuration time, not runtime. "
        "It writes files your agents already know how to read. "
        "No runtime dependency, no SDK, no vendor lock-in.",
    )
    add_heading(doc, "Core Principles", level=2)
    add_bullet(
        doc,
        "All configuration lives in .codi/ — a single, version-controlled directory.",
        "One source of truth.",
    )
    add_bullet(
        doc,
        "A change to .codi/ propagates to all agent config files via codi generate.",
        "Zero drift.",
    )
    add_bullet(
        doc,
        "Teams choose a preset; individuals add personal overrides in ~/.codi/user.yaml.",
        "Team + individual balance.",
    )
    add_bullet(
        doc,
        "Verification tokens detect when generated files are manually edited.",
        "Integrity verification.",
    )
    add_bullet(
        doc,
        "Flags can be locked to prevent overrides — security_scan: locked: true.",
        "Policy enforcement.",
    )
    add_bullet(
        doc,
        "Presets are shareable bundles — install from ZIP, GitHub, or registry.",
        "Reusable team configs.",
    )
    add_heading(doc, "How It Works — Four Steps", level=2)
    add_table(
        doc,
        ["Step", "Command", "What Happens"],
        [
            [
                "1. Initialize",
                "codi init",
                "Interactive wizard creates .codi/ with chosen preset",
            ],
            [
                "2. Customize",
                "codi add rule security",
                "Add rules, skills, agents to .codi/",
            ],
            [
                "3. Generate",
                "codi generate",
                "Write CLAUDE.md, .cursorrules, AGENTS.md, etc.",
            ],
            [
                "4. Monitor",
                "codi status",
                "Detect drift between .codi/ and generated files",
            ],
        ],
        col_widths=[1.2, 1.8, 3.5],
    )


def build_core_concepts(doc: Document) -> None:
    add_section_label(doc, "04", "Core Concepts")
    add_heading(doc, "The Five Artifact Types")
    add_body(
        doc,
        (
            "Codi organizes all configuration into five artifact types. "
            "Each type is a Markdown file with YAML frontmatter stored inside .codi/. "
            "Different agents support different artifact types — Codi's adapters handle "
            "the translation automatically."
        ),
    )
    add_heading(doc, "Rules", level=2)
    add_body(
        doc,
        (
            "Rules are instructions that AI agents follow during development sessions. "
            "They define coding standards, security requirements, workflow policies, "
            "and team conventions. Rules are the most portable artifact — every agent "
            "supports them."
        ),
    )
    add_table(
        doc,
        ["Rule", "Purpose"],
        [
            [
                "codi-security",
                "OWASP Top 10:2025, secret management, input validation, auth/authorization",
            ],
            [
                "codi-testing",
                "Testing strategy (integration > unit), 80% coverage, TDD, contract testing",
            ],
            [
                "codi-code-style",
                "Naming conventions, function size (<30 lines), immutability, no deep imports",
            ],
            [
                "codi-typescript",
                "strict: true, satisfies operator, path aliases, async patterns, Zod validation",
            ],
            [
                "codi-git-workflow",
                "Conventional commits, atomic commits, trunk-based dev, no force-push",
            ],
            [
                "codi-architecture",
                "Hexagonal architecture, modular monolith, dependency injection, CQRS",
            ],
            [
                "codi-performance",
                "Core Web Vitals, N+1 prevention, connection pooling, caching strategy",
            ],
            [
                "codi-documentation",
                "Diataxis framework, JSDoc, ADRs, runbooks, file naming convention",
            ],
            [
                "codi-error-handling",
                "Never swallow errors, Result types, circuit breakers, bulkhead isolation",
            ],
            [
                "codi-production-mindset",
                "OpenTelemetry, SLOs, feature flags, blue-green deployments",
            ],
        ],
        col_widths=[2.2, 4.3],
    )
    add_body(
        doc,
        "27 rule templates available. Each rule specifies: name, description, priority (high/medium/low), scope (file patterns), and managed_by (codi or user).",
        italic=True,
        color=RL3_GRAY,
    )
    add_heading(doc, "Skills", level=2)
    add_body(
        doc,
        (
            "Skills are reusable multi-step workflows that agents can invoke on demand. "
            "Unlike rules (always-on instructions), skills are activated by the user "
            "with a slash command or an explicit request. A skill can include scripts, "
            "reference documents, evaluation criteria, and sub-agents."
        ),
    )
    add_table(
        doc,
        ["Category", "Example Skills"],
        [
            [
                "Code Quality",
                "codi-code-review, codi-security-scan, codi-test-coverage, codi-e2e-testing",
            ],
            [
                "Git Workflow",
                "codi-commit (conventional commits, pre-commit checks, message validation)",
            ],
            [
                "Codi Platform",
                "codi-preset-creator, codi-rule-creator, codi-skill-creator, codi-agent-creator",
            ],
            [
                "Document Generation",
                "codi-doc-engine, codi-content-factory, codi-documentation",
            ],
            ["Design", "codi-frontend-design, codi-canvas-design, codi-rl3-brand"],
            [
                "Content",
                "codi-content-factory, codi-internal-comms, codi-codebase-onboarding",
            ],
        ],
        col_widths=[2.0, 4.5],
    )
    add_body(
        doc,
        "44 skill templates available. Skills live in .codi/skills/{name}/SKILL.md and support subdirectories: scripts/, references/, assets/, evals/.",
        italic=True,
        color=RL3_GRAY,
    )
    add_heading(doc, "Agents", level=2)
    add_body(
        doc,
        (
            "Agents are autonomous subagent definitions. When an AI coding assistant "
            "encounters a task that matches an agent's scope, it can delegate work to "
            "that agent with a dedicated context window, specific tools, and a defined process. "
            "Agents are supported by Claude Code and Codex."
        ),
    )
    add_table(
        doc,
        ["Agent", "Purpose", "Tools"],
        [
            [
                "codi-code-reviewer",
                "Senior code reviewer. Severity matrix (CRITICAL/HIGH/MEDIUM/LOW). Confidence filtering >80%.",
                "Read, Grep, Glob, Bash",
            ],
            [
                "codi-security-analyzer",
                "OWASP vulnerability scanner. Covers injection, auth, supply chain, API security.",
                "Read, Grep, Glob, Bash",
            ],
            [
                "codi-test-generator",
                "Generates unit, integration, and e2e test suites from existing code.",
                "Read, Write, Grep, Glob, Bash",
            ],
        ],
        col_widths=[1.8, 3.2, 1.5],
    )
    add_body(
        doc,
        "22 agent templates available. Agents are defined in .codi/agents/*.md.",
        italic=True,
        color=RL3_GRAY,
    )
    add_heading(doc, "Commands", level=2)
    add_body(
        doc,
        (
            "Commands are slash commands exposed in the AI agent interface. "
            "They provide one-word access to common workflows. "
            "Claude Code and Codex support slash commands natively."
        ),
    )
    add_table(
        doc,
        ["Command", "Purpose"],
        [
            [
                "/codi-commit",
                "Create a well-structured git commit following conventional commits format",
            ],
            [
                "/codi-review",
                "Review recent code changes for quality, security, and best practices",
            ],
            [
                "/codi-test-run",
                "Run the project test suite and report results with coverage",
            ],
            [
                "/codi-security-scan",
                "Run a comprehensive security scan and report vulnerabilities",
            ],
            [
                "/codi-test-coverage",
                "Measure test coverage and identify untested code paths",
            ],
            [
                "/codi-refactor",
                "Identify and safely remove dead code and redundant abstractions",
            ],
            ["/codi-onboard", "Analyze the codebase and generate an onboarding guide"],
        ],
        col_widths=[2.0, 4.5],
    )
    add_body(doc, "17 command templates available.", italic=True, color=RL3_GRAY)
    add_heading(doc, "Brands", level=2)
    add_body(
        doc,
        (
            "Brands are visual identity definitions — logos, color palettes, typography, "
            "tone of voice, and design guidelines. When a brand is installed, AI agents "
            "apply it when generating visual content, presentations, and documents. "
            "Brands live in .codi/brands/{name}/BRAND.md with an assets/ subdirectory."
        ),
    )


def build_flags(doc: Document) -> None:
    add_section_label(doc, "05", "The Flag System")
    add_heading(doc, "Behavioral Switches")
    add_body(
        doc,
        (
            "Flags are the behavioral control layer of Codi. "
            "They answer questions like: Can agents delete files? Must tests pass before a commit? "
            "Is force-push allowed? There are 16 flags covering security, testing, git safety, "
            "and language policies. Each flag has a mode and a value."
        ),
    )
    add_heading(doc, "Flag Modes", level=2)
    add_table(
        doc,
        ["Mode", "Behavior", "Override?"],
        [
            [
                "enforced",
                "Always active. Non-negotiable. Stops the resolution chain.",
                "No",
            ],
            ["locked", "Value cannot be changed by any higher layer.", "No"],
            ["enabled", "Active with the specified value.", "Yes"],
            ["disabled", "Explicitly turned off.", "Yes"],
            ["inherited", "Skip — use the parent layer's value.", "Yes"],
            [
                "conditional",
                "Apply only if agent, file pattern, or language matches.",
                "Yes",
            ],
        ],
        col_widths=[1.4, 3.8, 1.3],
    )
    add_heading(doc, "All 16 Flags", level=2)
    add_table(
        doc,
        ["Flag", "Type", "Default", "Purpose"],
        [
            [
                "auto_commit",
                "boolean",
                "false",
                "Automatic commits after agent changes",
            ],
            [
                "test_before_commit",
                "boolean",
                "true",
                "Pre-commit hook runs tests before allowing commit",
            ],
            [
                "security_scan",
                "boolean",
                "true",
                "Mandatory secret detection scan on staged files",
            ],
            [
                "type_checking",
                "enum",
                "strict",
                "Type checking level: strict / basic / off",
            ],
            ["require_tests", "boolean", "false", "Tests required for all new code"],
            ["allow_shell_commands", "boolean", "true", "Agent can run shell commands"],
            ["allow_file_deletion", "boolean", "true", "Agent can delete files"],
            ["lint_on_save", "boolean", "true", "Auto-lint files on save"],
            [
                "allow_force_push",
                "boolean",
                "false",
                "Allow force push to remote branches",
            ],
            ["require_pr_review", "boolean", "true", "PR review required before merge"],
            [
                "mcp_allowed_servers",
                "string[]",
                "[]",
                "Whitelist of allowed MCP server names",
            ],
            [
                "require_documentation",
                "boolean",
                "false",
                "Documentation required for all new code",
            ],
            [
                "allowed_languages",
                "string[]",
                "[*]",
                "Restrict which languages agents can use",
            ],
            [
                "progressive_loading",
                "enum",
                "metadata",
                "Skill inlining strategy for Windsurf/Cline",
            ],
            [
                "drift_detection",
                "enum",
                "warn",
                "How to report config drift: error / warn / off",
            ],
            [
                "auto_generate_on_change",
                "boolean",
                "true",
                "Auto-regenerate config when .codi/ changes",
            ],
        ],
        col_widths=[1.8, 0.9, 0.9, 2.9],
    )
    add_callout(
        doc,
        "Enforced flags in the recommended preset",
        "test_before_commit, security_scan, type_checking (strict), "
        "allow_force_push (false), and require_pr_review are all enforced — "
        "they cannot be overridden by personal user config or agent defaults.",
    )


def build_presets(doc: Document) -> None:
    add_section_label(doc, "06", "Presets")
    add_heading(doc, "One-Command Team Setup")
    add_body(
        doc,
        (
            "Presets are packaged bundles of flags and artifacts. "
            "They let teams share a complete AI agent configuration in a single ZIP file "
            "or GitHub repository. Install a preset with one command and your entire "
            ".codi/ directory is populated with curated rules, skills, agents, and policies."
        ),
    )
    add_heading(doc, "Built-In Presets", level=2)
    add_table(
        doc,
        ["Preset", "Target", "Key Characteristics"],
        [
            [
                "codi-minimal",
                "Rapid prototyping",
                "Security off, no test requirements, all actions permitted. Fast and flexible.",
            ],
            [
                "codi-balanced",
                "Most teams (default)",
                "Security on, strict type-checking, no force-push, PR review required.",
            ],
            [
                "codi-strict",
                "Regulated environments",
                "Security locked, tests required, file deletion restricted, no force-push.",
            ],
            [
                "codi-fullstack",
                "Web/app development",
                "Broad rule set: testing, security, TypeScript, architecture, documentation.",
            ],
            [
                "codi-dev",
                "Codi CLI development",
                "Strict TypeScript, anti-hardcoding, safe release workflow, full QA suite.",
            ],
            [
                "codi-power-user",
                "Advanced daily workflow",
                "Graph exploration, day tracking, error diagnosis, enhanced commit tools.",
            ],
        ],
        col_widths=[1.5, 1.8, 3.2],
    )
    add_heading(doc, "Preset Lifecycle", level=2)
    add_table(
        doc,
        ["Step", "Command", "Description"],
        [
            [
                "Create",
                "codi preset create my-setup",
                "Interactive wizard: choose flags, rules, skills, agents",
            ],
            ["Export", "codi preset export my-setup", "Package as .zip for sharing"],
            [
                "Install (ZIP)",
                "codi preset install ./my-setup.zip",
                "Apply preset to current project",
            ],
            [
                "Install (GitHub)",
                "codi preset install github:org/repo",
                "Install from public/private repository",
            ],
            ["List", "codi preset list", "Show installed and available presets"],
            [
                "Update",
                "codi preset update",
                "Pull newer versions of managed artifacts",
            ],
            [
                "Remove",
                "codi preset remove my-setup",
                "Clean up preset artifacts and state",
            ],
        ],
        col_widths=[1.2, 2.5, 2.8],
    )
    add_heading(doc, "Multi-Preset Repositories", level=2)
    add_body(
        doc,
        (
            "A single GitHub repository can contain multiple presets. "
            "Use the subfolder syntax to install a specific one: "
        ),
    )
    add_code_block(doc, "codi preset install github:my-org/codi-presets/backend-strict")
    add_callout(
        doc,
        "Managed vs user artifacts",
        "Preset-installed artifacts are marked managed_by: codi. "
        "They can be updated by codi preset update without overwriting your custom additions. "
        "Your own artifacts use managed_by: user and are never touched by Codi.",
    )
