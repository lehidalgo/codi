"""
generate_codi_docs.py

Generates a branded DOCX documentation for Codi using RL3 AI Agency brand guidelines.

Usage:
    /tmp/docx_venv/bin/python3 scripts/generate_codi_docs.py
"""

from datetime import datetime
from pathlib import Path

from docx import Document

from _doc_helpers import set_margins, setup_header_footer
from _doc_sections_a import (
    build_core_concepts,
    build_cover,
    build_executive_summary,
    build_flags,
    build_presets,
    build_problem,
    build_what_is_codi,
)
from _doc_sections_b import (
    build_agents_support,
    build_architecture,
    build_cli,
    build_verification,
    build_workflows,
)


def build_document() -> Path:
    doc = Document()

    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    set_margins(doc)
    setup_header_footer(doc)

    build_cover(doc)
    build_executive_summary(doc)
    doc.add_page_break()

    build_problem(doc)
    doc.add_page_break()

    build_what_is_codi(doc)
    doc.add_page_break()

    build_core_concepts(doc)
    doc.add_page_break()

    build_flags(doc)
    doc.add_page_break()

    build_presets(doc)
    doc.add_page_break()

    build_architecture(doc)
    doc.add_page_break()

    build_agents_support(doc)
    doc.add_page_break()

    build_cli(doc)
    doc.add_page_break()

    build_workflows(doc)
    doc.add_page_break()

    build_verification(doc)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{timestamp}_[GUIDE]_codi-documentation.docx"
    output_path = Path(__file__).parent.parent / "docs" / filename

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


if __name__ == "__main__":
    path = build_document()
    print(f"Generated: {path}")
