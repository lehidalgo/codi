"""
_doc_sections_b.py

Document sections 07-11: architecture, supported AI agents, CLI reference,
team workflows, verification and drift.
"""

from docx import Document

from _doc_helpers import (
    add_body,
    add_callout,
    add_code_block,
    add_heading,
    add_section_label,
    add_table,
)


def build_architecture(doc: Document) -> None:
    add_section_label(doc, "07", "Architecture")
    add_heading(doc, "How Codi Works Internally")
    add_heading(doc, "Three-Layer Configuration Resolution", level=2)
    add_body(doc, (
        "Codi resolves configuration by merging three layers in priority order. "
        "Higher layers override lower layers — except for flags marked as enforced or locked."
    ))
    add_table(doc,
        ["Layer", "Location", "Priority", "Override?"],
        [
            ["Preset", "Applied at install time", "Lowest", "Yes (unless locked)"],
            ["Repo", ".codi/ directory", "Middle", "Yes"],
            ["User", "~/.codi/user.yaml", "Highest", "Yes (enforced flags block this)"],
        ],
        col_widths=[1.2, 2.5, 1.2, 1.6]
    )
    add_body(doc, (
        "Important: presets are consumed at install time only. "
        "codi generate reads exclusively from .codi/ — the preset is not re-read on every generation. "
        "This makes generation fast, deterministic, and offline-capable."
    ))
    add_heading(doc, "Generation Pipeline", level=2)
    add_body(doc, "When you run codi generate, the following steps execute in sequence:")
    add_table(doc,
        ["Step", "Module", "Description"],
        [
            ["1. Scan", "config/parser.ts", "Read all .md files in .codi/, parse YAML frontmatter"],
            ["2. Validate", "config/validator.ts", "Check for duplicates, size limits, adapter existence"],
            ["3. Resolve", "config/resolver.ts", "Merge layers into a NormalizedConfig object"],
            ["4. Adapt", "adapters/*.ts", "Format config for each target agent"],
            ["5. Verify", "core/verify/", "Compute and inject verification token into each file"],
            ["6. Write", "core/generator/", "Write files to disk (or dry-run preview)"],
            ["7. Track", "core/audit/", "Update state.json with file hashes"],
        ],
        col_widths=[0.7, 1.8, 4.0]
    )
    add_heading(doc, "Adapter Pattern — Five Agents", level=2)
    add_body(doc, (
        "Each supported AI agent has a dedicated adapter that translates "
        "NormalizedConfig into that agent's native format. "
        "The adapter pattern means adding support for a new agent requires "
        "only writing a new adapter — zero changes to the core."
    ))
    add_table(doc,
        ["Adapter", "Instruction File", "MCP Config", "Notes"],
        [
            ["Claude Code", "CLAUDE.md", ".mcp.json", "Supports rules, skills, agents, commands"],
            ["Cursor", ".cursorrules", ".cursor/mcp.json", "Supports rules, skills. No agents or commands."],
            ["Codex", "AGENTS.md", ".codex/mcp.toml", "Supports rules, skills, agents"],
            ["Windsurf", ".windsurfrules", "—", "Rules and skills inline in single file"],
            ["Cline", ".clinerules", "—", "Rules and skills inline in single file"],
        ],
        col_widths=[1.3, 1.5, 1.5, 2.2]
    )
    add_heading(doc, "Verification & Drift Detection", level=2)
    add_body(doc, (
        "Every generated file contains a verification section appended by Codi after generation. "
        "The verification token is a 12-character hash of sorted artifact names — "
        "deterministic and stable across identical configurations."
    ))
    add_body(doc, (
        "When someone manually edits a generated file, the hash no longer matches. "
        "codi status detects this and reports the drift. "
        "The drift_detection flag controls behavior: warn (log only), error (exit 1, blocks CI), "
        "or off (no reporting)."
    ))


def build_agents_support(doc: Document) -> None:
    add_section_label(doc, "08", "Supported AI Agents")
    add_heading(doc, "Full Compatibility Matrix")
    add_table(doc,
        ["Feature", "Claude Code", "Cursor", "Codex", "Windsurf", "Cline"],
        [
            ["Instruction file", "CLAUDE.md", ".cursorrules", "AGENTS.md", ".windsurfrules", ".clinerules"],
            ["Rules", "✓", "✓", "✓", "✓", "✓"],
            ["Skills", "✓", "✓", "✓", "✓", "✓"],
            ["Agents", "✓", "—", "✓", "—", "—"],
            ["Slash commands", "✓", "—", "✓", "—", "—"],
            ["MCP servers", ".mcp.json", ".cursor/mcp.json", ".codex/mcp.toml", "—", "—"],
            ["Separate rule files", "✓", "✓", "—", "—", "—"],
            ["Separate skill dirs", "✓", "✓", "✓", "✓", "✓"],
        ],
        col_widths=[1.8, 1.1, 1.0, 1.0, 1.1, 1.0]
    )
    add_callout(doc, "Recommendation",
        "Start with Claude Code (most complete feature support) and add "
        "additional agents as needed. Codi handles the format differences automatically — "
        "your .codi/ configuration stays identical regardless of which agents you target."
    )


def build_cli(doc: Document) -> None:
    add_section_label(doc, "09", "CLI Reference")
    add_heading(doc, "All Commands")
    add_heading(doc, "Setup & Initialization", level=2)
    add_table(doc,
        ["Command", "Description"],
        [
            ["codi init", "Interactive wizard. Creates .codi/ with a selected preset. Supports --preset, --no-interactive."],
            ["codi add <type> <name>", "Add a rule, skill, agent, command, brand, or mcp-server. Use --template to choose a template."],
        ],
        col_widths=[2.5, 4.0]
    )
    add_heading(doc, "Generation & Output", level=2)
    add_table(doc,
        ["Command", "Description"],
        [
            ["codi generate", "Generate all agent config files from .codi/. Options: --agent, --dry-run, --force."],
            ["codi watch", "Watch .codi/ for changes and auto-regenerate on save."],
            ["codi clean", "Delete all generated config files (CLAUDE.md, .cursorrules, etc.)."],
        ],
        col_widths=[2.5, 4.0]
    )
    add_heading(doc, "Status & Health", level=2)
    add_table(doc,
        ["Command", "Description"],
        [
            ["codi status", "Show drift status between .codi/ source and generated files. --diff for colored unified diffs."],
            ["codi doctor", "Full health check: config valid, all files present, hooks installed."],
            ["codi validate", "Validate .codi/ configuration schema without generating."],
            ["codi verify", "Check verification tokens in generated files."],
            ["codi compliance", "Combined health + drift + verification check. CI-friendly, exits non-zero on issues."],
            ["codi ci", "CI-optimized version of compliance. Returns structured JSON output."],
        ],
        col_widths=[2.5, 4.0]
    )
    add_heading(doc, "Preset Management", level=2)
    add_table(doc,
        ["Command", "Description"],
        [
            ["codi preset create [name]", "Create a new preset with the interactive wizard."],
            ["codi preset list", "List all installed and built-in presets."],
            ["codi preset install <source>", "Install from ZIP file, GitHub repo, or registry."],
            ["codi preset export <name>", "Export a preset as a .zip file for sharing."],
            ["codi preset validate <name>", "Check preset structure and manifest."],
            ["codi preset remove <name>", "Remove a preset and its managed artifacts."],
            ["codi preset edit <name>", "Interactive editor for preset artifacts."],
            ["codi preset update", "Update all installed presets to latest versions."],
        ],
        col_widths=[2.5, 4.0]
    )
    add_heading(doc, "Maintenance", level=2)
    add_table(doc,
        ["Command", "Description"],
        [
            ["codi update", "Update managed artifacts (preset-managed rules, skills, etc.) in place."],
            ["codi revert", "Restore generated files from the last known good state backup."],
            ["codi skill export", "Export a skill for sharing or contribution."],
            ["codi contribute", "Guided workflow to share artifacts with the community."],
        ],
        col_widths=[2.5, 4.0]
    )
    add_heading(doc, "Global Options", level=2)
    add_table(doc,
        ["Option", "Effect"],
        [
            ["-j, --json", "Output results as JSON — useful for scripting and CI pipelines"],
            ["-v, --verbose", "Debug-level output including file paths and resolution details"],
            ["-q, --quiet", "Silent mode — only errors are printed"],
            ["--no-color", "Plain text output without ANSI color codes"],
        ],
        col_widths=[2.0, 4.5]
    )
    add_callout(doc, "Interactive Command Center",
        "Running codi with no arguments launches the interactive Command Center — "
        "a TUI with grouped actions, live status, and keyboard navigation. "
        "Ideal for first-time users and exploratory configuration."
    )


def build_workflows(doc: Document) -> None:
    add_section_label(doc, "10", "Team Workflows")
    add_heading(doc, "Getting Started in 4 Commands")
    add_code_block(doc, (
        "# 1. Install Codi globally\n"
        "npm install -g codi\n\n"
        "# 2. Initialize your project with the balanced preset\n"
        "codi init --preset codi-balanced\n\n"
        "# 3. Generate config files for all agents\n"
        "codi generate\n\n"
        "# 4. Verify everything is consistent\n"
        "codi doctor"
    ))
    add_heading(doc, "Daily Development Workflow", level=2)
    add_table(doc,
        ["Scenario", "What Codi Does"],
        [
            ["Add a new security rule", "codi add rule my-auth → edit .codi/rules/my-auth.md → codi generate"],
            ["Update team testing policy", "Edit .codi/rules/codi-testing.md → codi generate → commit .codi/ + generated files"],
            ["Check for config drift", "codi status → review diff → fix manually edited files or re-generate"],
            ["Onboard a new team member", "Git clone → codi generate → all agent configs instantly available"],
            ["Share preset with another team", "codi preset export my-setup → share .zip → they run codi preset install"],
        ],
        col_widths=[2.2, 4.3]
    )
    add_heading(doc, "CI/CD Integration", level=2)
    add_body(doc, "Add Codi compliance checks to your CI pipeline to catch drift before merge:")
    add_code_block(doc, (
        "# .github/workflows/codi.yml\n"
        "- name: Check Codi config consistency\n"
        "  run: npx codi compliance\n"
        "  # Exits non-zero if drift_detection: error and any file has drifted"
    ))
    add_heading(doc, "Pre-Commit Hooks", level=2)
    add_body(doc, (
        "Codi detects your existing hook runner (Husky, pre-commit, Lefthook, or standalone) "
        "and installs hooks adapted to that runner. Three core hooks are controlled by flags:"
    ))
    add_table(doc,
        ["Hook", "Triggered by Flag", "What It Does"],
        [
            ["Test runner", "test_before_commit: true", "Run test suite before allowing commit"],
            ["Secret detection", "security_scan: true", "Scan staged files for leaked secrets"],
            ["Type check", "type_checking: strict", "Run TypeScript compiler check on staged files"],
        ],
        col_widths=[1.8, 2.2, 2.5]
    )


def build_verification(doc: Document) -> None:
    add_section_label(doc, "11", "Verification & Drift")
    add_heading(doc, "Configuration Integrity")
    add_body(doc, (
        "Codi tracks the state of every generated file. "
        "After generation, each file contains a verification section with a token "
        "derived from the sorted list of artifact names. "
        "This token is deterministic — the same .codi/ configuration always produces "
        "the same token."
    ))
    add_body(doc, (
        "When someone manually edits CLAUDE.md or .cursorrules, the token no longer matches "
        "what Codi would generate. codi status detects this mismatch and reports it. "
        "codi verify inspects the token in isolation without regenerating."
    ))
    add_heading(doc, "Drift Detection Modes", level=2)
    add_table(doc,
        ["Mode", "Behavior", "Use Case"],
        [
            ["off", "No reporting — drift is silently ignored.", "Development environments with frequent manual edits"],
            ["warn", "Log warning, continue with exit code 0.", "Default — informational without blocking CI"],
            ["error", "Log error, exit with code 1.", "CI/CD pipelines — block merge if config has drifted"],
        ],
        col_widths=[0.9, 2.8, 2.8]
    )
    add_callout(doc, "Recommended configuration for teams",
        "Set drift_detection: error in your team preset so that any manually "
        "edited generated file blocks the pull request. "
        "All config changes must go through .codi/ — not by editing generated files directly."
    )
    add_heading(doc, "Managed vs User Artifacts", level=2)
    add_body(doc, (
        "The managed_by field on every artifact controls what Codi can and cannot update:"
    ))
    add_table(doc,
        ["managed_by", "Who Sets It", "Can Codi Update?", "Use For"],
        [
            ["codi", "Preset installation", "Yes — via codi preset update", "Team standards, framework rules"],
            ["user", "You (manually)", "Never", "Project-specific rules, personal preferences"],
        ],
        col_widths=[1.2, 1.8, 1.8, 1.7]
    )
