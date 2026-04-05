"""
_doc_helpers.py

RL3 brand colors, XML helpers, and style helpers for generate_codi_docs.py.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# RL3 Brand Colors
# ---------------------------------------------------------------------------
RL3_BLACK = RGBColor(0x0A, 0x0A, 0x0B)
RL3_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RL3_ACCENT = RGBColor(0xC8, 0xB8, 0x8A)   # Gold
RL3_GRAY = RGBColor(0x7A, 0x7A, 0x7A)
RL3_DARK_GRAY = RGBColor(0x1A, 0x1A, 0x1B)
RL3_MID_GRAY = RGBColor(0x2A, 0x2A, 0x2B)
RL3_LIGHT_BG = RGBColor(0xF5, 0xF5, 0xF5)
RL3_TABLE_ALT = RGBColor(0xFA, 0xF8, 0xF3)  # Warm near-white for alt rows


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, rgb: RGBColor) -> None:
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def set_cell_border(cell, side: str, size: int, color: str) -> None:
    """Add a border to one side of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color)
    tcBorders.append(border)


def add_horizontal_rule(doc: Document, color: RGBColor = RL3_ACCENT, thickness: int = 6):
    """Insert a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def add_section_label(doc: Document, number: str, title: str):
    """Gold monospace section label: '01 — Title'."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{number} — {title.upper()}")
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RL3_ACCENT
    run.font.bold = True
    add_horizontal_rule(doc, RL3_ACCENT, 4)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    """Bold near-black heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.color.rgb = RL3_BLACK
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(16)
    else:
        run.font.size = Pt(13)
    return p


def add_body(doc: Document, text: str, italic: bool = False, color: RGBColor | None = None):
    """Standard body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.italic = italic
    run.font.color.rgb = color if color else RL3_DARK_GRAY
    return p


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None):
    """Bullet point with optional bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = RL3_BLACK
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RL3_DARK_GRAY
    return p


def add_callout(doc: Document, title: str, text: str):
    """Gold-accented callout box using a 1-column table."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, RL3_TABLE_ALT)
    set_cell_border(cell, "left", 16, f"{RL3_ACCENT[0]:02X}{RL3_ACCENT[1]:02X}{RL3_ACCENT[2]:02X}")
    title_p = cell.paragraphs[0]
    title_p.paragraph_format.space_before = Pt(4)
    title_p.paragraph_format.space_after = Pt(2)
    tr = title_p.add_run(title)
    tr.font.name = "Calibri"
    tr.font.size = Pt(11)
    tr.font.bold = True
    tr.font.color.rgb = RL3_ACCENT
    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(4)
    br = body_p.add_run(text)
    br.font.name = "Calibri"
    br.font.size = Pt(10.5)
    br.font.color.rgb = RL3_DARK_GRAY
    doc.add_paragraph()
    return table


def add_table(doc: Document, headers: list, rows: list, col_widths: list | None = None):
    """Styled table with dark header row and alternating body rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, RL3_BLACK)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(h)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RL3_WHITE

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = RL3_TABLE_ALT if r_idx % 2 == 0 else RL3_WHITE
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if isinstance(val, tuple):
                run = p.add_run(val[0])
                run.font.bold = val[1]
            else:
                run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = RL3_DARK_GRAY

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_code_block(doc: Document, code: str) -> None:
    """Monospace code block with light background."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, RGBColor(0xF0, 0xEE, 0xE8))
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1B)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Header / Footer / Margins
# ---------------------------------------------------------------------------

def setup_header_footer(doc: Document) -> None:
    """Configure RL3 page header and footer for all sections."""
    section = doc.sections[0]

    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()

    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.paragraph_format.space_after = Pt(0)

    r1 = hp.add_run("RL")
    r1.font.name = "Calibri"
    r1.font.size = Pt(18)
    r1.font.bold = True
    r1.font.color.rgb = RL3_BLACK

    r2 = hp.add_run("3")
    r2.font.name = "Calibri"
    r2.font.size = Pt(18)
    r2.font.bold = True
    r2.font.color.rgb = RL3_ACCENT

    r3 = hp.add_run("  AI AGENCY")
    r3.font.name = "Courier New"
    r3.font.size = Pt(8)
    r3.font.color.rgb = RL3_GRAY

    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{RL3_ACCENT[0]:02X}{RL3_ACCENT[1]:02X}{RL3_ACCENT[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)

    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()

    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    fpPr = fp._p.get_or_add_pPr()
    fpBdr = OxmlElement("w:pBdr")
    top_b = OxmlElement("w:top")
    top_b.set(qn("w:val"), "single")
    top_b.set(qn("w:sz"), "4")
    top_b.set(qn("w:space"), "1")
    top_b.set(qn("w:color"), f"{RL3_ACCENT[0]:02X}{RL3_ACCENT[1]:02X}{RL3_ACCENT[2]:02X}")
    fpBdr.append(top_b)
    fpPr.append(fpBdr)

    frun = fp.add_run("RL3 AI AGENCY  |  codi  |  ")
    frun.font.name = "Courier New"
    frun.font.size = Pt(8)
    frun.font.color.rgb = RL3_GRAY

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Courier New")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "16")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{RL3_GRAY[0]:02X}{RL3_GRAY[1]:02X}{RL3_GRAY[2]:02X}")
    rPr.extend([rFonts, sz, color_el])
    run_el.append(rPr)
    run_el.append(fldChar1)
    run_el.append(instrText)
    run_el.append(fldChar2)
    fp._p.append(run_el)


def set_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
