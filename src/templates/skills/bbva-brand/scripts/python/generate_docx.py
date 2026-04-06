#!/usr/bin/env python3
"""
generate_docx.py — BBVA-branded DOCX generator (python-docx, FALLBACK runtime).

Usage:
    python3 generate_docx.py --content content.json --output output.docx

Install dependency: pip install python-docx
content.json schema: {title, subtitle?, author?, sections: [{number, label, heading, body, items?, callout?}]}
"""

import sys
import json
import argparse
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent))
import brand_tokens as bt

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _rgb(key: str) -> RGBColor:
    r, g, b = bt.rgb(key)
    return RGBColor(r, g, b)


def _set_font(
    run,
    color_key: str,
    size_pt: int,
    font_key: str = "pptx_body",
    bold=False,
    italic=False,
) -> None:
    run.font.color.rgb = _rgb(color_key)
    run.font.size = Pt(size_pt)
    run.font.name = bt.FONTS.get(font_key, "Arial")
    run.font.bold = bold
    run.font.italic = italic


def _add_callout(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_font(run, "primary_dark", 12, italic=True)

    # Add left border via direct XML (python-docx lacks native paragraph borders)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), bt.hex_color("primary").upper())
    pBdr.append(left)
    pPr.append(pBdr)

    # Shading
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "E8EBF7")
    pPr.append(shd)

    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.left_indent = Cm(0.5)


def build_cover(doc: Document, content: dict) -> None:
    # BBVA wordmark
    p = doc.add_paragraph()
    run = p.add_run("BBVA")
    _set_font(run, "primary", 18, bold=True)
    p.paragraph_format.space_after = Pt(20)

    # Title
    p = doc.add_paragraph()
    run = p.add_run(content["title"])
    _set_font(run, "primary_dark", 28, font_key="pptx_headlines", bold=True)
    p.paragraph_format.space_after = Pt(12)

    # Subtitle
    if content.get("subtitle"):
        p = doc.add_paragraph()
        run = p.add_run(content["subtitle"])
        _set_font(run, "text_secondary", 14, italic=True)
        p.paragraph_format.space_after = Pt(24)

    # Author + date
    meta_parts = []
    if content.get("author"):
        meta_parts.append(content["author"])
    meta_parts.append(datetime.now().strftime("%B %Y"))
    p = doc.add_paragraph()
    run = p.add_run(" · ".join(meta_parts))
    _set_font(run, "text_light", 11)
    p.paragraph_format.space_after = Pt(36)

    doc.add_page_break()


def build_section(doc: Document, section: dict) -> None:
    # Section number + label (heading 2)
    p = doc.add_paragraph()
    run = p.add_run(f"{section['number']}  {section['label'].upper()}")
    _set_font(run, "primary", 11, bold=True)
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(6)

    # Horizontal rule (via bottom border on paragraph)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), bt.hex_color("primary").upper())
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Section heading (heading 1)
    heading = doc.add_heading(section["heading"], level=1)
    for run in heading.runs:
        _set_font(run, "primary_dark", 20, font_key="pptx_headlines", bold=True)
    heading.paragraph_format.space_after = Pt(12)

    # Body
    p = doc.add_paragraph(section["body"])
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(18)
    for run in p.runs:
        _set_font(run, "text_primary", 12)

    # Bullet items
    for item in section.get("items", []):
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            _set_font(run, "text_primary", 12)
        p.paragraph_format.space_after = Pt(4)

    # Callout
    if section.get("callout"):
        _add_callout(doc, section["callout"])


def generate_docx(content: dict, output_path: str) -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    build_cover(doc, content)

    for sec in content.get("sections", []):
        build_section(doc, sec)

    doc.save(output_path)
    n = len(content.get("sections", []))
    print(f"DOCX written: {output_path} ({n} sections)")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="BBVA DOCX generator (python-docx fallback)"
    )
    parser.add_argument("--content", required=True, help="Path to content.json")
    parser.add_argument("--output", default="output.docx", help="Output .docx file")
    args = parser.parse_args()

    with open(args.content, "r", encoding="utf-8") as f:
        content = json.load(f)

    generate_docx(content, args.output)


if __name__ == "__main__":
    main()
