"""
generate_docx.py — Generates an RL3-branded DOCX document from structured content.

Usage:
    python generate_docx.py --content content.json --output output.docx
"""

import sys
import os
import argparse
import json
import zipfile
import io
import re

sys.path.insert(0, os.path.dirname(__file__))
import brand_tokens as bt

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


GOLD = RGBColor(0xC8, 0xB8, 0x8A)
FIXED_TIMESTAMP = "2026-01-01T00:00:00Z"


# ── Validation ──────────────────────────────────────────────────────


def _validate_content(content: dict) -> None:
    """Raise ValueError if required fields are missing."""
    if "title" not in content:
        raise ValueError("Missing required field: 'title'")
    if "sections" not in content:
        raise ValueError("Missing required field: 'sections'")
    if not isinstance(content["sections"], list) or len(content["sections"]) < 1:
        raise ValueError("'sections' must be a non-empty list")
    for i, sec in enumerate(content["sections"]):
        for key in ("number", "label", "heading", "body"):
            if key not in sec:
                raise ValueError(f"Section {i} missing required field: '{key}'")


# ── Document Generation ─────────────────────────────────────────────


def generate_docx(content: dict, output_path: str) -> None:
    """Generate an RL3-branded DOCX from *content* and save to *output_path*."""
    _validate_content(content)

    doc = Document()

    # ── Header: "RL" + "3" (gold) ────────────────────────────────
    heading = doc.add_heading(level=1)
    run_rl = heading.add_run("RL")
    run_rl.bold = True
    run_rl.font.name = bt.FONTS["headlines"]
    run_rl.font.size = Pt(28)

    run_3 = heading.add_run("3")
    run_3.bold = True
    run_3.font.name = bt.FONTS["headlines"]
    run_3.font.size = Pt(28)
    run_3.font.color.rgb = GOLD

    # Subtitle / title line
    if content.get("subtitle"):
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run(f"{content['title']} | {content['subtitle']}")
        run_sub.font.name = bt.FONTS["headlines"]
        run_sub.font.size = Pt(14)
    else:
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run(content["title"])
        run_sub.font.name = bt.FONTS["headlines"]
        run_sub.font.size = Pt(14)

    # Date
    if content.get("date"):
        p_date = doc.add_paragraph()
        run_date = p_date.add_run(content["date"])
        run_date.font.name = bt.FONTS["body"]
        run_date.font.size = Pt(10)
        run_date.font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)

    # Gold horizontal rule
    hr = doc.add_paragraph()
    pPr = hr._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="C8B88A"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    # ── Sections ─────────────────────────────────────────────────
    for sec in content["sections"]:
        # Section heading: "01 — Contexto: Analisis inicial"
        sec_heading = doc.add_heading(level=2)
        sec_text = f"{sec['number']} \u2014 {sec['label']}: {sec['heading']}"
        run_h = sec_heading.add_run(sec_text)
        run_h.bold = True
        run_h.font.name = bt.FONTS["headlines"]
        run_h.font.size = Pt(16)

        # Body paragraph
        p_body = doc.add_paragraph()
        run_body = p_body.add_run(sec["body"])
        run_body.font.name = bt.FONTS["body"]
        run_body.font.size = Pt(11)

        # Callout (optional)
        if sec.get("callout"):
            p_call = doc.add_paragraph()
            run_call = p_call.add_run(sec["callout"])
            run_call.italic = True
            run_call.font.name = bt.FONTS["body"]
            run_call.font.size = Pt(11)
            run_call.font.color.rgb = GOLD
            # Left indent for visual distinction
            p_call.paragraph_format.left_indent = Inches(0.5)

        # Items (optional)
        if sec.get("items"):
            for item in sec["items"]:
                p_item = doc.add_paragraph()
                run_item = p_item.add_run(f"\u2192 {item}")
                run_item.font.name = bt.FONTS["body"]
                run_item.font.size = Pt(11)
                p_item.paragraph_format.left_indent = Inches(0.3)

    # ── Footer ───────────────────────────────────────────────────
    footer_text = content.get("footer_text", "RL3 AI AGENCY")
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    run_footer = footer_para.add_run(footer_text)
    run_footer.font.name = bt.FONTS["headlines"]
    run_footer.font.size = Pt(8)
    run_footer.font.color.rgb = RGBColor(0x7A, 0x7A, 0x7A)

    # ── Save ─────────────────────────────────────────────────────
    doc.save(output_path)

    # ── Strip timestamps for determinism ─────────────────────────
    _strip_timestamps(output_path)


def _strip_timestamps(path: str) -> None:
    """Rewrite the docx ZIP so that dcterms:created and dcterms:modified
    are set to a fixed timestamp, ensuring SHA-256 determinism."""
    buf = io.BytesIO()
    with zipfile.ZipFile(path, "r") as zin:
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "docProps/core.xml":
                    text = data.decode("utf-8")
                    # Replace created timestamp
                    text = re.sub(
                        r"(<dcterms:created[^>]*>)[^<]*(</dcterms:created>)",
                        rf"\g<1>{FIXED_TIMESTAMP}\g<2>",
                        text,
                    )
                    # Replace modified timestamp
                    text = re.sub(
                        r"(<dcterms:modified[^>]*>)[^<]*(</dcterms:modified>)",
                        rf"\g<1>{FIXED_TIMESTAMP}\g<2>",
                        text,
                    )
                    data = text.encode("utf-8")
                # Normalize zip entry metadata for determinism
                item.date_time = (2026, 1, 1, 0, 0, 0)
                item.compress_type = zipfile.ZIP_DEFLATED
                zout.writestr(item, data)

    with open(path, "wb") as f:
        f.write(buf.getvalue())


# ── CLI ──────────────────────────────────────────────────────────────


def main():
    parser = argparse.ArgumentParser(
        description="Generate an RL3-branded DOCX document."
    )
    parser.add_argument(
        "--content",
        required=True,
        help="Path to a JSON file with document content",
    )
    parser.add_argument(
        "--output",
        required=True,
        help="Output path for the .docx file",
    )
    args = parser.parse_args()

    if not os.path.isfile(args.content):
        print(f"Error: content file not found: {args.content}", file=sys.stderr)
        sys.exit(1)

    with open(args.content, "r", encoding="utf-8") as f:
        content = json.load(f)

    generate_docx(content, args.output)
    print(f"Generated: {args.output}")


if __name__ == "__main__":
    main()
