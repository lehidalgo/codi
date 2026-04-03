r"""
doc_validator.py — Validates a DOCX file against RL3 brand rules.

5 rules:
  1. heading_bold    — At least one heading/run with bold=True
  2. section_numbers — At least one paragraph matching regex 0\d (01, 02, 03 format)
  3. gold_accent     — At least one run with font.color.rgb == RGBColor(0xc8, 0xb8, 0x8a)
  4. footer_text     — Footer contains "RL3 AI AGENCY"
  5. forbidden_phrases — No phrase from brand_tokens.PHRASES_AVOID in any paragraph text

Usage:
    python doc_validator.py --input file.docx
"""

import sys, os, re, argparse, json

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
import brand_tokens as bt

from docx import Document
from docx.shared import RGBColor


GOLD = RGBColor(0xC8, 0xB8, 0x8A)


def validate_docx(filepath: str) -> dict:
    """Validate a DOCX file against RL3 brand rules.

    Returns dict with keys: passed (bool), errors (list), warnings (list).
    """
    doc = Document(filepath)
    errors: list[dict] = []
    warnings: list[dict] = []

    # ── Rule 1: heading_bold ────────────────────────────────────────
    has_bold_heading = False
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            for run in para.runs:
                if run.bold:
                    has_bold_heading = True
                    break
        if has_bold_heading:
            break
    if not has_bold_heading:
        errors.append({
            "rule": "heading_bold",
            "message": "No heading run with bold=True found.",
        })

    # ── Rule 2: section_numbers ─────────────────────────────────────
    section_pattern = re.compile(r"0\d")
    has_section_number = False
    for para in doc.paragraphs:
        if section_pattern.search(para.text):
            has_section_number = True
            break
    if not has_section_number:
        errors.append({
            "rule": "section_numbers",
            "message": "No paragraph with section number format (01, 02, 03...) found.",
        })

    # ── Rule 3: gold_accent ─────────────────────────────────────────
    has_gold = False
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.color.rgb == GOLD:
                has_gold = True
                break
        if has_gold:
            break
    if not has_gold:
        errors.append({
            "rule": "gold_accent",
            "message": f"No run with gold accent color ({GOLD}) found.",
        })

    # ── Rule 4: footer_text ─────────────────────────────────────────
    has_footer = False
    for section in doc.sections:
        footer = section.footer
        for para in footer.paragraphs:
            if "RL3 AI AGENCY" in para.text:
                has_footer = True
                break
        if has_footer:
            break
    if not has_footer:
        errors.append({
            "rule": "footer_text",
            "message": 'Footer does not contain "RL3 AI AGENCY".',
        })

    # ── Rule 5: forbidden_phrases ───────────────────────────────────
    found_phrases: list[str] = []
    for para in doc.paragraphs:
        text_lower = para.text.lower()
        for phrase in bt.PHRASES_AVOID:
            if phrase.lower() in text_lower:
                found_phrases.append(phrase)
    if found_phrases:
        errors.append({
            "rule": "forbidden_phrases",
            "message": f"Forbidden phrases found: {', '.join(sorted(set(found_phrases)))}",
        })

    return {
        "passed": len(errors) == 0,
        "errors": errors,
        "warnings": warnings,
    }


# ── CLI ─────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Validate a DOCX file against RL3 brand rules."
    )
    parser.add_argument("--input", required=True, help="Path to the .docx file")
    args = parser.parse_args()

    if not os.path.isfile(args.input):
        print(f"Error: file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    result = validate_docx(args.input)
    print(json.dumps(result, indent=2, ensure_ascii=False))
    sys.exit(0 if result["passed"] else 1)


if __name__ == "__main__":
    main()
